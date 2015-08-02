# coding: utf-8
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
DIR = os.path.dirname(os.path.abspath(__file__))

def name_of_pic(name):
    return os.path.join(DIR, name)

document = Document()

head=document.add_heading('Document Title', 0)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

names = os.listdir(DIR)
for name in names:
    if name.endswith('png') or name.endswith('PNG'):
        picture=document.add_picture(name_of_pic(name), width=Inches(1.25))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p=document.add_paragraph(' Add a comment to this picture ')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.save(sys.argv[1])
